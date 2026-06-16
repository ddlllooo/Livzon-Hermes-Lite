"""文档导入管线：解析 → 分块 → Embedding → 写入 OceanBase。

支持格式：PDF / DOCX / Markdown / 纯文本
分块策略：父块（完整段落/章节）→ 子块（滑动窗口，用于检索）
"""

import hashlib
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)


# ── 数据结构 ─────────────────────────────────────────────────────────

@dataclass
class Document:
    """原始文档。"""
    content: str
    source: str         # 文件名/URL/标识
    file_type: str      # pdf/docx/md/txt
    metadata: dict = field(default_factory=dict)


@dataclass
class ParentChunk:
    """父块 — 完整语义单元，用于 LLM Context。"""
    parent_id: str
    content: str
    source: str
    metadata: dict = field(default_factory=dict)


@dataclass
class ChildChunk:
    """子块 — 检索单元，带 Embedding。"""
    chunk_id: str
    parent_id: str
    content: str
    chunk_idx: int
    embedding: List[float] = field(default_factory=list)


# ── 文档解析 ─────────────────────────────────────────────────────────

def parse_file(file_path: str) -> Document:
    """解析文件为 Document 对象。"""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = path.suffix.lower()
    source = path.name

    if suffix == ".pdf":
        return _parse_pdf(path, source)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(path, source)
    elif suffix in (".md", ".markdown"):
        return _parse_text(path, source, "md")
    elif suffix in (".txt", ".text", ".csv", ".json", ".jsonl"):
        return _parse_text(path, source, "txt")
    else:
        # 尝试当纯文本读取
        return _parse_text(path, source, "txt")


def parse_text_content(content: str, source: str = "inline",
                       file_type: str = "txt") -> Document:
    """从文本内容直接创建 Document（用于 API 上传）。"""
    return Document(content=content, source=source,
                    file_type=file_type, metadata={})


def _parse_pdf(path: Path, source: str) -> Document:
    """解析 PDF — 优先 pymupdf，备选 marker-pdf。"""
    try:
        import pymupdf  # PyMuPDF
        doc = pymupdf.open(str(path))
        pages = []
        for page in doc:
            pages.append(page.get_text())
        content = "\n\n".join(pages)
        doc.close()
        return Document(content=content, source=source, file_type="pdf",
                        metadata={"pages": len(pages)})
    except ImportError:
        pass

    try:
        from marker.converters.pdf import PdfConverter
        converter = PdfConverter()
        content = converter(str(path))
        return Document(content=content, source=source, file_type="pdf",
                        metadata={})
    except ImportError:
        raise RuntimeError(
            "PDF 解析需要 pymupdf: pip install pymupdf\n"
            "或 marker-pdf: pip install marker-pdf"
        )


def _parse_docx(path: Path, source: str) -> Document:
    """解析 DOCX。"""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        return Document(content=content, source=source, file_type="docx",
                        metadata={"paragraphs": len(paragraphs)})
    except ImportError:
        raise RuntimeError("DOCX 解析需要 python-docx: pip install python-docx")


def _parse_text(path: Path, source: str, file_type: str) -> Document:
    """解析纯文本/Markdown。"""
    # 检测编码
    raw = path.read_bytes()
    encoding = _detect_encoding(raw)
    content = raw.decode(encoding, errors="replace")
    return Document(content=content, source=source, file_type=file_type,
                    metadata={"encoding": encoding, "size_bytes": len(raw)})


def _detect_encoding(data: bytes) -> str:
    """检测文本编码。"""
    try:
        import chardet
        result = chardet.detect(data[:10000])
        return result.get("encoding") or "utf-8"
    except ImportError:
        pass
    # 回退：尝试 utf-8，失败用 gbk
    try:
        data[:10000].decode("utf-8")
        return "utf-8"
    except UnicodeDecodeError:
        return "gbk"


# ── 分块策略 ─────────────────────────────────────────────────────────

def split_into_parent_chunks(
    doc: Document,
    max_parent_chars: int = 2000,
    overlap_chars: int = 200,
) -> List[ParentChunk]:
    """将文档切分为父块。

    策略：
      - Markdown：按标题（#）分段
      - 通用文本：按段落（双换行）分段，过长段落按字符数切割
    """
    source_id = _make_id(doc.source)

    if doc.file_type == "md":
        sections = _split_markdown(doc.content)
    else:
        sections = _split_by_paragraphs(doc.content)

    # 合并过短段落，拆分过长段落
    merged = _merge_and_split(sections, max_parent_chars, overlap_chars)

    parents = []
    for i, text in enumerate(merged):
        text = text.strip()
        if not text:
            continue
        pid = f"{source_id}_p{i:04d}"
        parents.append(ParentChunk(
            parent_id=pid,
            content=text,
            source=doc.source,
            metadata={**doc.metadata, "chunk_index": i},
        ))

    return parents


def split_parent_into_children(
    parent: ParentChunk,
    child_chars: int = 512,
    child_overlap: int = 64,
) -> List[ChildChunk]:
    """将父块切分为子块（滑动窗口）。

    子块用于检索命中的最小单元，命中后反查父块完整文本。
    """
    text = parent.content
    if len(text) <= child_chars:
        cid = f"{parent.parent_id}_c0000"
        return [ChildChunk(
            chunk_id=cid,
            parent_id=parent.parent_id,
            content=text,
            chunk_idx=0,
        )]

    children = []
    start = 0
    idx = 0
    while start < len(text):
        end = start + child_chars
        # 尽量在句号/换行处断开
        if end < len(text):
            break_pos = _find_break(text, end)
            if break_pos > start:
                end = break_pos
        chunk_text = text[start:end].strip()
        if chunk_text:
            cid = f"{parent.parent_id}_c{idx:04d}"
            children.append(ChildChunk(
                chunk_id=cid,
                parent_id=parent.parent_id,
                content=chunk_text,
                chunk_idx=idx,
            ))
            idx += 1
        start = end - child_overlap if end < len(text) else end

    return children


# ── 文档处理管线 ─────────────────────────────────────────────────────

def process_document(
    doc: Document,
    embedder,
    chunk_store,
    parent_chars: int = 2000,
    child_chars: int = 512,
    child_overlap: int = 64,
) -> dict:
    """完整文档处理管线：解析 → 分块 → Embedding → 写入 OceanBase。

    返回统计信息。
    """
    # ① 切分父块
    parents = split_into_parent_chunks(doc, max_parent_chars=parent_chars)
    logger.info("文档 '%s': %d 个父块", doc.source, len(parents))

    # ② 切分子块
    all_children: List[ChildChunk] = []
    for parent in parents:
        children = split_parent_into_children(
            parent, child_chars=child_chars, child_overlap=child_overlap
        )
        all_children.extend(children)

    logger.info("文档 '%s': %d 个子块", doc.source, len(all_children))

    # ③ 批量生成 Embedding
    texts = [c.content for c in all_children]
    embeddings = embedder.embed_batch(texts)
    for child, emb in zip(all_children, embeddings):
        child.embedding = emb

    # ④ 写入 OceanBase
    # 先删除旧数据（如果同一文档重复导入）
    if parents:
        chunk_store.delete_document(parents[0].parent_id.rsplit("_p", 1)[0])

    # 写入父块
    for parent in parents:
        chunk_store.insert_parent(
            parent_id=parent.parent_id,
            content=parent.content,
            source=parent.source,
            metadata=parent.metadata,
        )

    # 批量写入子块
    child_rows = [
        {
            "chunk_id": c.chunk_id,
            "parent_id": c.parent_id,
            "content": c.content,
            "embedding": c.embedding,
        }
        for c in all_children
    ]
    inserted = chunk_store.insert_children(child_rows)

    stats = {
        "source": doc.source,
        "file_type": doc.file_type,
        "total_chars": len(doc.content),
        "parent_chunks": len(parents),
        "child_chunks": len(all_children),
        "inserted": inserted,
    }
    logger.info("导入完成: %s", stats)
    return stats


def process_file(
    file_path: str,
    embedder,
    chunk_store,
    parent_chars: int = 2000,
    child_chars: int = 512,
    child_overlap: int = 64,
) -> dict:
    """从文件路径导入：解析 → 分块 → Embedding → 写入。"""
    doc = parse_file(file_path)
    return process_document(
        doc, embedder, chunk_store,
        parent_chars=parent_chars,
        child_chars=child_chars,
        child_overlap=child_overlap,
    )


def process_directory(
    dir_path: str,
    embedder,
    chunk_store,
    extensions: List[str] = None,
    **kwargs,
) -> List[dict]:
    """批量导入目录下的所有文件。"""
    if extensions is None:
        extensions = [".pdf", ".docx", ".doc", ".md", ".txt", ".text"]

    path = Path(dir_path)
    if not path.is_dir():
        raise NotADirectoryError(f"目录不存在: {dir_path}")

    results = []
    files = sorted(
        f for f in path.rglob("*")
        if f.is_file() and f.suffix.lower() in extensions
    )
    logger.info("扫描到 %d 个文件: %s", len(files), dir_path)

    for f in files:
        try:
            stat = process_file(str(f), embedder, chunk_store, **kwargs)
            results.append(stat)
        except Exception as e:
            logger.error("处理文件失败 '%s': %s", f.name, e)
            results.append({"source": f.name, "error": str(e)})

    return results


# ── 内部工具函数 ─────────────────────────────────────────────────────

def _make_id(text: str) -> str:
    """生成短哈希 ID。"""
    return hashlib.md5(text.encode()).hexdigest()[:12]


def _split_markdown(content: str) -> List[str]:
    """按 Markdown 标题分段。"""
    sections = []
    current = []
    for line in content.split("\n"):
        if re.match(r"^#{1,4}\s+", line) and current:
            sections.append("\n".join(current))
            current = [line]
        else:
            current.append(line)
    if current:
        sections.append("\n".join(current))
    return sections


def _split_by_paragraphs(content: str) -> List[str]:
    """按双换行分段。"""
    paragraphs = re.split(r"\n\s*\n", content)
    return [p.strip() for p in paragraphs if p.strip()]


def _merge_and_split(
    sections: List[str],
    max_chars: int,
    overlap: int,
) -> List[str]:
    """合并过短段落，拆分过长段落。"""
    result = []
    buffer = ""

    for section in sections:
        if len(section) > max_chars:
            # 先 flush buffer
            if buffer:
                result.append(buffer)
                buffer = ""
            # 拆分过长段落
            start = 0
            while start < len(section):
                end = min(start + max_chars, len(section))
                result.append(section[start:end])
                start = end - overlap if end < len(section) else end
        elif len(buffer) + len(section) + 2 > max_chars:
            result.append(buffer)
            buffer = section
        else:
            buffer = f"{buffer}\n\n{section}" if buffer else section

    if buffer:
        result.append(buffer)

    return result


def _find_break(text: str, pos: int) -> int:
    """在 pos 附近寻找合适的断句位置（句号/换行）。"""
    search_range = min(100, pos // 4)
    for offset in range(search_range):
        for p in (pos - offset, pos + offset):
            if 0 <= p < len(text) and text[p] in "。\n.!?\n":
                return p + 1
    return pos
